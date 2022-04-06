import re
import datetime
from docx import Document
from docx.shared import Inches


def main():
    paragraph_data = {}
    table_data = {}
    template = 'temp.docx'
    result = 'result.docx'
    counter = 0
    with open("data.txt", "r", encoding='utf-8') as f:
        lines = [line.rstrip('\n') for line in f]
        for i in range(4):
            paragraph_data[lines[i]] = lines[i + 4]
        paragraph_data['{{DATE}}'] = str(datetime.datetime.now().date())
        print(paragraph_data['{{DATE}}'])
        for line in lines:
            match = re.findall(r'\d{2}.\d{2}.\d{4}', line)
            if (len(match) > 1):
                counter += 1
                table_data[counter] = {'START_DATE': match[0], 'FINISH_DATE': match[1], 'DURATION': line.split(' ')[2]}

    template_doc = Document(template)


    for key, value in paragraph_data.items():
        for paragraph in template_doc.paragraphs:
            replace_text(paragraph, key, value)

    for table in template_doc.tables:
        counter = 1
        for key, value in table_data.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(counter)
            row_cells[1].text = value['START_DATE']
            row_cells[2].text = value['FINISH_DATE']
            row_cells[3].text = value['DURATION']
            counter += 1

    template_doc.save(result)


def replace_text(paragraph, key, value):
    if key in paragraph.text:
        paragraph.text = paragraph.text.replace(key, value)


main()
