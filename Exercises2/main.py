from openpyxl import load_workbook
from docx import Document
import datetime


def main():
    company_data = {}
    data = {}
    number = 1
    template = 'temp.docx'
    result = 'result.docx'
    wb = load_workbook('data_isukk_8.xlsx', data_only=True)
    ws = wb.sheetnames
    for sheet in ws:
        for row in wb[sheet].iter_rows(min_row=2, max_row=wb[sheet].max_row, values_only=True):
            company_data[row[1]] = [row[0], row[2]]
            data[row[1]] = []

        for sender, receiver, reason, product, unit, price, amount, summ in wb[sheet]:
            for key, value in data.items():
                if receiver.value == key:
                    value.append([product.value, unit.value, price.value, amount.value, summ.value])

    print(data)
    for key, value in data.items():
        template_doc = Document(template)
        total_items = 0
        total_sum = 0
        for table in template_doc.tables:
            counter = 1
            for lists in value:
                counter2 = 0
                row_cells = table.add_row().cells
                row_cells[0].text = str(counter)
                counter += 1
                for val in lists:
                    if counter2 == 3:
                        total_items += val
                    elif counter2 == 4:
                        total_sum += val
                    counter2 += 1
                    row_cells[counter2].text = str(val)


        for key1, value1 in company_data.items():
            basic_inf = {}
            basic_inf['{{NUMBER}}'] = number
            basic_inf['{{DATE}}'] = str(datetime.datetime.now().date())
            basic_inf['{{SENDER}}'] = value1[0]
            basic_inf['{{RECEIVER}}'] = key1
            basic_inf['{{REASON}}'] = value1[1]
            basic_inf['{{TOTALNUMBEROFPRODUCTS}}'] = total_items
            basic_inf['{{TOTALSUM}}'] = total_sum
            for key2, value2 in basic_inf.items():
                 for paragraph in template_doc.paragraphs:
                    replace_text(paragraph, key2, str(value2))

        template_doc.save(f'накладная_{number}_{key}.docx')
        number += 1


def replace_text(paragraph, key, value):
    if key in paragraph.text:
        paragraph.text = paragraph.text.replace(key, value)


main()
